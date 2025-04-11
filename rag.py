import os
from typing import List, Dict, Any, Optional, Tuple
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchResults
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pdfplumber
from docx import Document as DocxDocument
from tqdm import tqdm
import re


class SimpleRAG:
    def __init__(self, google_api_key: str):
        self.google_api_key = google_api_key
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash", 
            google_api_key=google_api_key,
            temperature=0.1
        )
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        self.search_tool = DuckDuckGoSearchResults()
        self.vectorstore = None
        self.current_document = None  # Store current document text
        self.current_file_name = None  # Store current file name
        self.current_summary = None  # Store current document summary

    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX files."""
        if file_path.lower().endswith('.pdf'):
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in tqdm(pdf.pages, desc="Processing PDF"):
                    text += page.extract_text() or ""
            return text
        
        elif file_path.lower().endswith('.docx'):
            doc = DocxDocument(file_path)
            return "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file_path.lower().endswith(('.txt', '.md')):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    def add_document(self, file_path: str) -> None:
        """Add a document to the RAG system, replacing any existing document."""
        try:
            # Clear existing document data
            self.current_document = None
            self.current_file_name = None
            self.current_summary = None
            self.vectorstore = None
            
            # Extract text from document
            text = self.extract_text(file_path)
            
            # Store the current document info
            self.current_file_name = os.path.basename(file_path)
            self.current_document = text
            
            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=400,
                separators=["\n\n", "\n", ". ", " "]
            )
            chunks = text_splitter.split_text(text)
            
            # Create documents and update vector store
            docs = [Document(page_content=chunk, metadata={"source": file_path}) for chunk in chunks]
            self.vectorstore = FAISS.from_documents(docs, self.embeddings)
                
            print(f"Successfully loaded document: {file_path}")
            
            # Auto-generate summary if it looks like a research paper
            if self.is_likely_research_paper(text):
                print("This appears to be a research paper. Generating summary...")
                self.current_summary = self.summarize_paper(content=text)
                print(f"Summary generated for {self.current_file_name}")
            
        except Exception as e:
            print(f"Error adding document {file_path}: {e}")

    def is_likely_research_paper(self, text: str) -> bool:
        """Determine if text is likely a research paper."""
        academic_indicators = [
            r'\babstract\b', r'\bintroduction\b', r'\bmethodology\b', 
            r'\bresults\b', r'\bdiscussion\b', r'\bconclusion\b',
            r'\breferences\b', r'\bcitation\b', r'\bfigure \d+\b'
        ]
        
        # Count how many academic indicators are present
        indicator_count = sum(1 for pattern in academic_indicators if re.search(pattern, text, re.IGNORECASE))
        
        # If at least 3 indicators are present, likely a research paper
        return indicator_count >= 3
    
    def search_web(self, query: str) -> List[Dict[str, str]]:
        """Search the web using DuckDuckGo."""
        try:
            result = self.search_tool.run(query)
            return [{"title": "DuckDuckGo Search", "content": result}]
        except Exception as e:
            print(f"Web search error: {e}")
            return []
        
    def answer_question(self, query: str, use_web_search: bool = True) -> Dict[str, Any]:
        """Answer questions using current document knowledge and web search."""
        if not self.current_document:
            return {
                "answer": "No document is currently loaded. Please add a document first.",
                "sources": []
            }

        context_parts = []
        sources = set()  # Changed to set to avoid duplicates
        
        # Always add summary context first if available
        if self.current_summary:
            summary_context = f"""DOCUMENT SUMMARY:
Title: {self.current_summary['title']}
Field: {self.current_summary['paper_type']}
Authors: {self.current_summary['authors']}

ABSTRACT:
{self.current_summary['abstract']}

KEY FINDINGS:
{self.current_summary['key_findings']}

METHODOLOGY:
{self.current_summary['methodology']}

RESEARCH GAPS AND FUTURE DIRECTIONS:
{self.current_summary['gaps_and_future']}

CONCLUSION:
{self.current_summary['conclusion']}"""
            context_parts.append(summary_context)
            sources.add(self.current_file_name)
        
        # Analyze question type
        question_analysis_prompt = f"""Analyze this research-related question to determine its type and requirements:
Question: {query}

Determine:
1. Question type (e.g., factual, analytical, comparative, methodological)
2. Required information sources (paper content, summaries, methodology, results)
3. Scope of answer needed (brief fact, detailed explanation, analysis)
"""
        question_analysis = self.llm.invoke(question_analysis_prompt).content

        # Get relevant documents from current document
        if self.vectorstore is not None:
            relevant_docs = self.vectorstore.similarity_search(query, k=4)
            for doc in relevant_docs:
                context_parts.append(doc.page_content)
                sources.add(self.current_file_name)

        # Add web search results if enabled
        if use_web_search:
            web_results = self.search_web(query)
            for result in web_results:
                context_parts.append(result["content"])
                sources.add("DuckDuckGo Search")

        # Create enhanced prompt with context
        prompt = f"""Based on the following information, provide a comprehensive answer to the question. 
Consider both the high-level summary and detailed content.
If you cannot answer based on the provided information, clearly state so.
Always cite your sources and specify which paper or document you're referring to.

QUESTION ANALYSIS:
{question_analysis}

CONTEXT (including summary and relevant sections):
{' '.join(context_parts)}

Question: {query}

Please provide:
1. A direct answer to the question
2. Supporting evidence from both the summary and detailed content
3. Any relevant context or caveats
4. Citations for specific claims

Answer:"""

        # Generate answer
        answer = self.llm.invoke(prompt)
        
        return {
            "answer": answer.content,
            "sources": list(sources),
            "question_analysis": question_analysis
        }
    
    def detect_paper_type(self, text: str) -> str:
        """Detect the type of research paper based on content analysis."""
        # Check for specific patterns to determine paper type
        prompt = f"""Analyze the following research paper text and determine its primary field or discipline (e.g., computer science, medicine, physics, etc.):

First 3000 characters of the paper:
{text[:3000]}

Based on the text, what is the most likely academic field or discipline of this research paper?"""

        response = self.llm.invoke(prompt)
        return response.content.strip()
    
    def extract_paper_sections(self, text: str) -> Dict[str, str]:
        """Extract common sections from research papers."""
        # Common section headers in research papers
        common_sections = [
            "abstract", "introduction", "background", "related work", 
            "methodology", "methods", "experimental setup", "experiments",  
            "results", "discussion", "conclusion", "future work", "references"
        ]
        
        sections = {}
        
        # Try to find section headers and extract content
        for i, section in enumerate(common_sections):
            # Create regex pattern for the section header
            pattern = re.compile(f"(?i)\\b{section}\\b[\\s:.]*", re.IGNORECASE)
            matches = list(pattern.finditer(text))
            
            if matches:
                start_idx = matches[0].end()
                
                # Find next section if it exists
                end_idx = None
                for next_section in common_sections[i+1:]:
                    next_pattern = re.compile(f"(?i)\\b{next_section}\\b[\\s:.]*", re.IGNORECASE)
                    next_matches = list(next_pattern.finditer(text))
                    if next_matches:
                        end_idx = next_matches[0].start()
                        break
                
                # If no end index found, use the rest of the text (up to a reasonable limit)
                if end_idx is None:
                    end_idx = min(start_idx + 10000, len(text))
                
                # Extract section content
                sections[section] = text[start_idx:end_idx].strip()
        
        return sections
    
    def extract_key_info(self, text: str) -> Dict[str, str]:
        """Extract key information like title, authors, and publication date."""
        # Get first 1000 characters which typically contain metadata
        header_text = text[:1000]
        
        prompt = f"""Extract the following information from this research paper header:
1. Title
2. Authors
3. Publication date/year
4. Affiliation/Institution
5. DOI (if present)

Text:
{header_text}

Provide the information in this format:
Title: [paper title]
Authors: [list of authors]
Publication: [date/year]
Affiliation: [institutions]
DOI: [doi if present, otherwise "Not specified"]"""

        response = self.llm.invoke(prompt)
        
        # Parse the response into a dictionary
        info = {}
        for line in response.content.strip().split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                info[key.strip().lower()] = value.strip()
        
        return info
    
    def summarize_paper(self, file_name: str = None, content: str = None) -> Dict[str, Any]:
        """Summarize a research paper with comprehensive analysis."""
        if file_name is None and content is None:
            return {"error": "Must provide either file_name or content"}
        
        if file_name is not None:
            if file_name not in self.current_file_name:
                return {"error": f"File {file_name} not found in loaded documents"}
            content = self.current_document
        
        # Detect paper discipline
        paper_type = self.detect_paper_type(content)
        
        # Extract basic paper information
        paper_info = self.extract_key_info(content)
        
        # Extract sections
        sections = self.extract_paper_sections(content)
        
        # Generate overall summary with enhanced prompt
        summary_prompt = f"""Create a detailed and comprehensive summary of this research paper. Analyze the following aspects:

Abstract: {sections.get('abstract', 'Not available')}

Introduction: {sections.get('introduction', 'Not available')}

Methodology: {sections.get('methodology', sections.get('methods', 'Not available'))}

Results: {sections.get('results', 'Not available')}

Discussion/Conclusion: {sections.get('discussion', sections.get('conclusion', 'Not available'))}

Please provide a thorough analysis covering:
1. Research Context and Gap:
   - What is the broader context of this research?
   - What gap or problem does it address?

2. Research Objectives:
   - What are the main research questions?
   - What are the specific goals?

3. Methodological Approach:
   - What methods were used?
   - Why were these methods chosen?
   - What is the research design?

4. Key Findings:
   - What are the main results?
   - How do they relate to the research questions?

5. Significance and Impact:
   - What are the theoretical implications?
   - What are the practical applications?
   - How does this advance the field?

6. Limitations and Future Work:
   - What are the main limitations?
   - What future research directions are suggested?

7. Innovation and Contribution:
   - What is novel about this work?
   - How does it contribute to the field?

Write this as a well-structured, cohesive summary that thoroughly captures the paper's essence and contribution to the field.
"""
        
        summary = self.llm.invoke(summary_prompt).content
        
        # Extract key findings with enhanced prompt
        findings_prompt = f"""Based on the paper's content, provide a comprehensive analysis of the key findings and implications:

Results section: {sections.get('results', 'Not available')}

Discussion: {sections.get('discussion', 'Not available')}

Conclusion: {sections.get('conclusion', 'Not available')}

Please analyze and present:
1. Major Findings (3-5 points):
   - List each significant finding
   - Explain its importance
   - Connect it to research objectives

2. Practical Implications:
   - Real-world applications
   - Industry relevance
   - Potential impact

3. Theoretical Contributions:
   - Advances in knowledge
   - Theoretical framework contributions
   - Relationship to existing theories

Format each finding as a clear, detailed statement with supporting context.
"""
        
        key_findings = self.llm.invoke(findings_prompt).content
        
        # Extract research gaps and future directions
        gaps_prompt = f"""Analyze the paper's discussion of research gaps and future directions:

Introduction: {sections.get('introduction', 'Not available')}
Discussion: {sections.get('discussion', 'Not available')}
Conclusion: {sections.get('conclusion', 'Not available')}

Please identify:
1. Current gaps in the field
2. Limitations of existing approaches
3. Suggested future research directions
4. Potential improvements and extensions
"""
        
        gaps_and_future = self.llm.invoke(gaps_prompt).content
        
        # Create enhanced summary object
        paper_summary = {
            "title": paper_info.get('title', 'Unknown Title'),
            "authors": paper_info.get('authors', 'Unknown Authors'),
            "publication": paper_info.get('publication', 'Unknown Publication Date'),
            "paper_type": paper_type,
            "summary": summary,
            "key_findings": key_findings,
            "gaps_and_future": gaps_and_future,
            "abstract": sections.get('abstract', 'Abstract not available'),
            "methodology": sections.get('methodology', sections.get('methods', 'Methodology not available')),
            "conclusion": sections.get('conclusion', 'Conclusion not available'),
            "doi": paper_info.get('doi', 'Not specified'),
            "affiliation": paper_info.get('affiliation', 'Not specified')
        }
        
        # Store the summary if file_name was provided
        if file_name is not None:
            self.current_summary = paper_summary
            
        return paper_summary
    
    def format_paper_summary(self, summary: Dict[str, Any]) -> str:
        """Format the paper summary for display."""
        formatted = f"""
=== RESEARCH PAPER SUMMARY ===

TITLE: {summary['title']}
AUTHORS: {summary['authors']}
PUBLICATION: {summary['publication']}
FIELD: {summary['paper_type']}
AFFILIATION: {summary['affiliation']}
DOI: {summary['doi']}

ABSTRACT:
{summary['abstract']}

COMPREHENSIVE SUMMARY:
{summary['summary']}

KEY FINDINGS AND IMPLICATIONS:
{summary['key_findings']}

RESEARCH GAPS AND FUTURE DIRECTIONS:
{summary['gaps_and_future']}

METHODOLOGY OVERVIEW:
{summary['methodology']}

CONCLUSION:
{summary['conclusion']}
"""
        return formatted

    def save_summary_as_docx(self, output_path: str = None) -> None:
        """Save the current document summary as a DOCX file."""
        if not self.current_summary:
            print("No summary available. Please generate a summary first using 'summarize' command.")
            return

        if not output_path:
            # Create default filename based on original document name
            base_name = os.path.splitext(self.current_file_name)[0]
            output_path = f"{base_name}_summary.docx"

        try:
            doc = DocxDocument()
            
            # Add title
            doc.add_heading(self.current_summary['title'], 0)
            
            # Add metadata section
            doc.add_heading('Document Information', level=1)
            doc.add_paragraph(f"Original File: {self.current_file_name}")
            doc.add_paragraph(f"Authors: {self.current_summary['authors']}")
            doc.add_paragraph(f"Field: {self.current_summary['paper_type']}")
            if self.current_summary.get('doi'):
                doc.add_paragraph(f"DOI: {self.current_summary['doi']}")
            if self.current_summary.get('affiliation'):
                doc.add_paragraph(f"Affiliation: {self.current_summary['affiliation']}")
            
            # Add abstract
            doc.add_heading('Abstract', level=1)
            doc.add_paragraph(self.current_summary['abstract'])
            
            # Add comprehensive summary
            doc.add_heading('Comprehensive Summary', level=1)
            doc.add_paragraph(self.current_summary['summary'])
            
            # Add key findings
            doc.add_heading('Key Findings and Implications', level=1)
            doc.add_paragraph(self.current_summary['key_findings'])
            
            # Add methodology
            doc.add_heading('Methodology Overview', level=1)
            doc.add_paragraph(self.current_summary['methodology'])
            
            # Add research gaps and future directions
            doc.add_heading('Research Gaps and Future Directions', level=1)
            doc.add_paragraph(self.current_summary['gaps_and_future'])
            
            # Add conclusion
            doc.add_heading('Conclusion', level=1)
            doc.add_paragraph(self.current_summary['conclusion'])
            
            # Save the document
            doc.save(output_path)
            print(f"Summary successfully saved to: {output_path}")
            
        except Exception as e:
            print(f"Error saving summary: {e}")


def interactive_qa():
    """Run an interactive QA session."""
    google_api_key = os.environ.get("GOOGLE_API_KEY", "AIzaSyAb0h0btQBSzpwInYUhy2YZA9k7U2Zc510")
    rag = SimpleRAG(google_api_key)
    
    print("\n=== Research Paper Analysis & QA System ===")
    print("\nCommands:")
    print("1. add file: [path]          - Load a new document")
    print("2. summarize                 - Generate comprehensive paper summary")
    print("3. current file              - Show currently loaded document")
    print("4. save summary [path]       - Save summary as DOCX (path optional)")
    print("5. exit                      - Exit the system")
    print("\nFor questions, simply type your query.")
    
    while True:
        try:
            user_input = input("\nYou: ").strip()
            
            if user_input.lower() == 'exit':
                print("Thank you for using the Research Paper Analysis System. Goodbye!")
                break
                
            elif user_input.lower() == 'current file':
                if not rag.current_file_name:
                    print("No document currently loaded. Use 'add file: [path]' to load a document.")
                else:
                    print(f"\nCurrently loaded document: {rag.current_file_name}")
                continue
                
            elif user_input.lower().startswith('add file:'):
                file_path = user_input[9:].strip()
                if not file_path:
                    print("Please provide a file path. Usage: add file: [path]")
                    continue
                    
                print(f"\nProcessing file: {file_path}")
                print("This may take a moment...")
                rag.add_document(file_path)
                continue
                
            elif user_input.lower() == 'summarize':
                if not rag.current_document:
                    print("No document loaded. Please add a document first using 'add file: [path]'")
                    continue
                
                if not rag.current_summary:
                    print("Generating comprehensive summary...")
                    rag.current_summary = rag.summarize_paper(content=rag.current_document)
                
                print(rag.format_paper_summary(rag.current_summary))
                continue

            elif user_input.lower().startswith('save summary'):
                # Extract optional path from command
                parts = user_input.split(' ', 2)
                output_path = parts[2].strip() if len(parts) > 2 else None
                rag.save_summary_as_docx(output_path)
                continue
            
            # Handle regular questions
            print("\nAnalyzing question and searching for relevant information...")
            result = rag.answer_question(user_input)
            
            print("\nQuestion Analysis:")
            print("=" * 50)
            print(result["question_analysis"])
            
            print("\nAnswer:")
            print("=" * 50)
            print(result["answer"])
            
            if result["sources"]:
                print("\nSources Used:")
                print("=" * 50)
                for source in result["sources"]:
                    print(f"- {source}")
                    
        except Exception as e:
            print(f"\nError: {str(e)}")
            print("Please try again or use 'exit' to quit.")

if __name__ == "__main__":
    interactive_qa()